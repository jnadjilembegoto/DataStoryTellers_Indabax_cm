import streamlit as st
import joblib
import numpy as np
from docx import Document
from io import BytesIO

from Datas.data_link import data_dir
path = data_dir('modele_eligibilite.joblib')

def ml_analyse():
    # Charger le modèle
    model = joblib.load(path)
   # Interface utilisateur avec Streamlit
    st.title("🤖 Prédiction d'Éligibilité au Don de Sang")
    st.markdown("---")
    st.write("Cochez les reponses applicables sur le potentiel donneur")
    st.markdown("---")
    # Collecte des données utilisateur
    genre = st.radio("Genre", ["Homme", "Femme"])
    age = st.number_input("Âge", min_value=0, max_value=120, value=30)
    taux_hemoglobine = st.number_input("Taux d’hémoglobine", min_value=0.0, max_value=20.0, value=13.5)
    
    donne_sang = st.checkbox("A déjà donné le sang ?")
    antibiotique = st.checkbox("Sous antibiothérapie ?")
    if taux_hemoglobine<13:
        hemoglobine_basse=1
    else:
        hemoglobine_basse=0
    dernier_don_3mois = st.checkbox("Dernier don dans les 3 mois ?")
    ist_recente = st.checkbox("IST récente ?")
    
    # Variables spécifiques aux femmes
    if genre == "Femme":
        allaitement = st.checkbox("Allaitement en cours ?")
        accouchement_6mois = st.checkbox("Accouchement dans les 6 mois ?")
        ivg_6mois = st.checkbox("IVG dans les 6 mois ?")
        enceinte = st.checkbox("Enceinte ?")
        ddr_mauvais = st.checkbox("Dernières règles < 14 jours ?")
    else:
        allaitement = accouchement_6mois = ivg_6mois = enceinte = ddr_mauvais = False
    
    transfusion_antecedent = st.checkbox("Antécédents de transfusion ?")
    porteur_maladies = st.checkbox("Porteur de maladies ?")
    opere = st.checkbox("Opéré récemment ?")
    drepanocytaire = st.checkbox("Drepanocytaire ?")
    diabetique = st.checkbox("Diabétique ?")
    hypertendu = st.checkbox("Hypertendu ?")
    asthmatique = st.checkbox("Asthmatique ?")
    cardiaque = st.checkbox("Cardiaque ?")
    tatoue = st.checkbox("Tatoué ?")
    scarifie = st.checkbox("Scarifié ?")
    
    # Transformation des valeurs en format numérique
    donnees = np.array([age,
                        1 if genre == "Homme" else 0,
                        int(donne_sang),
                        taux_hemoglobine,
                        int(antibiotique),
                        int(hemoglobine_basse),
                        int(dernier_don_3mois),
                        int(ist_recente),
                        int(ddr_mauvais),
                        int(allaitement),
                        int(accouchement_6mois),
                        int(ivg_6mois),
                        int(enceinte),
                        int(transfusion_antecedent),
                        int(porteur_maladies),
                        int(opere),
                        int(drepanocytaire),
                        int(diabetique),
                        int(hypertendu),
                        int(asthmatique),
                        int(cardiaque),
                        int(tatoue),
                        int(scarifie)
                        ]).reshape(1, -1)
    st.markdown("---")
    st.write("Cliquez sur le bouton suivant pour prédire l'état d'éligibilité du potentiel donneur")
    st.markdown("---")
    # Bouton de prédiction
    if st.button("Prédire l'éligibilité"):
        prediction = model.predict(donnees)[0]
        
        # Résultats de la prédiction
        if prediction == 1:
            st.success("✅ L'individu est Eligible.")
            status = "Eligible"
        elif prediction == 2:
            st.error("❌ L'individu est Définitivement non-eligible.")
            status = "Définitivement non-eligible"
        elif prediction == 3:
            st.warning("⚠️ L'individu est Temporairement non-eligible.")
            status = "Temporairement non-eligible"
        else:
            st.info(f"ℹ️ Modalité inconnue : {prediction}")
            status = "Modalité inconnue"
        
        # Affichage du bouton de téléchargement seulement après la prédiction
        if 'show_report_button' not in st.session_state:
            st.session_state.show_report_button = True

        # Créer un rapport Word
        if st.session_state.show_report_button:
            document = Document()
            document.add_heading('Rapport d\'Éligibilité au Don de Sang', 0)
            document.add_paragraph(f"Genre: {genre}")
            document.add_paragraph(f"Âge: {age}")
            document.add_paragraph(f"Taux d’hémoglobine: {taux_hemoglobine}")
            document.add_paragraph(f"A déjà donné le sang: {'Oui' if donne_sang else 'Non'}")
            document.add_paragraph(f"Sous antibiothérapie: {'Oui' if antibiotique else 'Non'}")
            document.add_paragraph(f"Hémoglobine basse: {'Oui' if hemoglobine_basse else 'Non'}")
            document.add_paragraph(f"Dernier don dans les 3 mois: {'Oui' if dernier_don_3mois else 'Non'}")
            document.add_paragraph(f"IST récente: {'Oui' if ist_recente else 'Non'}")

            # Variables spécifiques aux femmes
            if genre == "Femme":
                document.add_paragraph(f"Allaitement en cours: {'Oui' if allaitement else 'Non'}")
                document.add_paragraph(f"Accouchement dans les 6 mois: {'Oui' if accouchement_6mois else 'Non'}")
                document.add_paragraph(f"IVG dans les 6 mois: {'Oui' if ivg_6mois else 'Non'}")
                document.add_paragraph(f"Enceinte: {'Oui' if enceinte else 'Non'}")
                document.add_paragraph(f"Dernières règles < 14 jours: {'Oui' if ddr_mauvais else 'Non'}")

            document.add_paragraph(f"Antécédents de transfusion: {'Oui' if transfusion_antecedent else 'Non'}")
            document.add_paragraph(f"Porteur de maladies: {'Oui' if porteur_maladies else 'Non'}")
            document.add_paragraph(f"Opéré récemment: {'Oui' if opere else 'Non'}")
            document.add_paragraph(f"Drepanocytaire: {'Oui' if drepanocytaire else 'Non'}")
            document.add_paragraph(f"Diabétique: {'Oui' if diabetique else 'Non'}")
            document.add_paragraph(f"Hypertendu: {'Oui' if hypertendu else 'Non'}")
            document.add_paragraph(f"Asthmatique: {'Oui' if asthmatique else 'Non'}")
            document.add_paragraph(f"Cardiaque: {'Oui' if cardiaque else 'Non'}")
            document.add_paragraph(f"Tatoué: {'Oui' if tatoue else 'Non'}")
            document.add_paragraph(f"Scarifié: {'Oui' if scarifie else 'Non'}")

            document.add_paragraph(f"\nStatut de l'éligibilité: {status}")

            
            # Sauvegarder le document en mémoire pour le télécharger
            byte_io = BytesIO()
            document.save(byte_io)
            byte_io.seek(0)
            st.markdown("---")
            st.write("Cochez les reponses applicables sur le potentiel donneur")
            st.markdown("---")
            # Bouton pour télécharger le rapport
            st.download_button(
                label="Télécharger le rapport",
                data=byte_io,
                file_name="rapport_eligibilite.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )



